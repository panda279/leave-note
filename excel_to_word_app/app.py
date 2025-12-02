import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# 设置页面标题
st.title("📊 Excel转Word工具 (学院精确排序版)")
st.write("自动清理空格后，按指定顺序严格排序学院数据")

# 第一步：上传Excel文件
st.header("第一步：上传Excel文件")
excel_file = st.file_uploader("选择Excel文件", type=['xlsx'])

# 定义学院排序顺序
COLLEGE_ORDER = [
    "经济与管理学院",
    "法学院",
    "文学与传媒学院", 
    "数据科学与人工智能学院",
    "电子与电气学院",
    "机器人学院",
    "建筑与能源工程学院",
    "设计艺术学院",
    "外国语学院",
    "创新创业学院"
]

def create_word_document(df, selected_columns):
    """创建Word文档并确保字体统一，添加请假说明和落款"""
    # 创建文档
    doc = Document()
    
    # 设置文档默认字体（确保所有文本统一）
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    
    # --- 新增：添加请假说明文字 ---
    # 标题
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run('各二级学院：')
    title_run.font.name = '宋体'
    title_run.font.size = Pt(12)
    title_run.font.bold = True  # 加粗
    
    # 正文
    text_content = """兹定于X年X月X日举办“XXX（填活动名称）”活动。以下同学因参与活动组织工作，将于X月X日 上午/下午/全天（根据实际时间选择）协助相关会务工作，无法参加该时间段课程。
特此申请为以下同学办理 X月X日 上午/下午/全天 的公假手续，恳请贵学院予以批准，谢谢！"""
    
    text_paragraph = doc.add_paragraph()
    text_run = text_paragraph.add_run(text_content)
    text_run.font.name = '宋体'
    text_run.font.size = Pt(10.5)
    
    # 添加一个空行分隔
    doc.add_paragraph()
    # --- 请假说明结束 ---
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(selected_columns))
    
    # 设置表格样式（可选，让表格更好看）
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    for i, col in enumerate(selected_columns):
        # 清空单元格内容
        header_cells[i].text = ''
        paragraph = header_cells[i].paragraphs[0]
        
        # 添加文本并设置字体
        run = paragraph.add_run(str(col))
        run.font.name = '宋体'
        run.font.size = Pt(11)
        run.font.bold = True  # 表头加粗
        
        # 居中对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加数据行
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(selected_columns):
            value = row[col]
            
            # 清空单元格内容
            row_cells[i].text = ''
            paragraph = row_cells[i].paragraphs[0]
            
            # 添加文本并设置字体
            text_content = str(value) if pd.notna(value) else ""
            run = paragraph.add_run(text_content)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            
            # 左对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # --- 新增：添加表格后的落款信息 ---
    # 添加一个空行
    doc.add_paragraph()
    
    # 创建落款段落（右对齐）
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 第一行：共青团温州理工学院委员会
    run1 = signature_paragraph.add_run('共青团温州理工学院委员会')
    run1.font.name = '宋体'
    run1.font.size = Pt(10.5)
    run1.font.bold = True  # 单位名称加粗
    
    # 添加换行
    signature_paragraph.add_run('\n')
    
    # 第二行：日期
    run2 = signature_paragraph.add_run('xx年xx月xx日')
    run2.font.name = '宋体'
    run2.font.size = Pt(10.5)
    # --- 落款结束 ---
    
    return doc

if excel_file is not None:
    # 读取Excel文件
    df = pd.read_excel(excel_file)
    
    # 显示原始数据预览
    st.subheader("数据预览 (原始)")
    st.write(f"总共有 {len(df)} 行数据")
    st.dataframe(df)
    
    # 第二步：检查并处理"学院"列
    st.header("第二步：处理学院排序")
    
    # 检查是否存在"学院"列
    if '学院' not in df.columns:
        st.error("错误：在Excel文件中未找到名为'学院'的列。请检查列名。")
        st.stop()
    
    # 核心步骤1：自动删除空格
    st.info("正在清理'学院'列中的空格...")
    df['学院'] = df['学院'].astype(str).str.strip()
    
    # 显示清理后的唯一值
    unique_colleges = df['学院'].unique()
    st.write("**清理空格后，'学院'列的唯一值有：**", unique_colleges.tolist())
    
    # 核心步骤2：按指定顺序重组数据
    st.info("正在按指定顺序重组数据...")
    
    # 创建一个空的DataFrame来存放排序后的结果
    sorted_dfs = []
    
    # 按照指定顺序，逐个学院提取数据
    for college in COLLEGE_ORDER:
        # 筛选出当前学院的行
        college_data = df[df['学院'] == college]
        if not college_data.empty:
            sorted_dfs.append(college_data)
            st.write(f"  ✓ 已提取: {college} ({len(college_data)}行)")
        else:
            st.write(f"  ⚠ 未找到: {college} (0行)")
    
    # 合并所有排序后的数据
    if sorted_dfs:
        df_sorted = pd.concat(sorted_dfs, ignore_index=True)
        
        # 处理不在指定顺序中的其他学院
        other_colleges = set(df['学院'].unique()) - set(COLLEGE_ORDER)
        if other_colleges:
            st.warning(f"发现以下未在排序列表中的学院，它们将被放在最后：{list(other_colleges)}")
            other_data = df[df['学院'].isin(other_colleges)]
            df_sorted = pd.concat([df_sorted, other_data], ignore_index=True)
        
        # 显示排序后的数据
        st.subheader("数据预览 (按学院排序后)")
        st.dataframe(df_sorted)
        
        # 更新df为排序后的数据
        df = df_sorted
    else:
        st.error("未匹配到任何指定学院的数据。请检查'学院'列的值。")
        st.stop()
    
    # 第三步：选择列
    st.header("第三步：选择要导出的列")
    all_columns = df.columns.tolist()
    selected_columns = st.multiselect(
        "选择要添加到Word的列",
        all_columns,
        default=all_columns[:4] if len(all_columns) >= 4 else all_columns
    )
    
    # 第四步：生成Word文档
    st.header("第四步：生成Word文档")
    
    if st.button("生成Word文档") and selected_columns:
        with st.spinner("正在生成Word文档..."):
            # 创建Word文档
            doc = create_word_document(df, selected_columns)
            
            # 保存到内存
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            # 提供下载
            st.success("文档生成成功！")
            st.download_button(
                label="📥 下载Word文档",
                data=file_stream,
                file_name="按学院排序的表格.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

else:
    st.info("请先上传Excel文件")
