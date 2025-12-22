import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
from datetime import datetime

# 设置页面标题
st.title("📄 请假单生成工具")
st.write("自动处理Excel数据，按学院排序生成格式规范的请假单")

# 上传Excel文件
excel_file = st.file_uploader("选择Excel文件", type=['xlsx', 'xls'])

# 定义学院排序顺序
COLLEGE_ORDER = [
    "经济与管理学院",
    "法学院",
    "文学与传媒学院", 
    "数据科学与人工智能学院",
    "建筑与能源工程学院",
    "电子与电气工程学院",
    "机器人工程学院",
    "设计艺术学院",
    "外国语学院",
    "创新创业学院"
]

def create_word_document(df, selected_columns, doc_type, activity_info):
    """创建Word文档，支持三种请假单类型"""
    # 创建文档
    doc = Document()
    
    # 设置全局字体
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.size = Pt(0)
    
    # 文档大标题
    title_paragraph = doc.add_paragraph()
    if doc_type == "公假单":
        title_text = '公假单'
    elif doc_type == "抵晚单":
        title_text = '抵晚单'
    else:  # 早自习请假单
        title_text = '早自习请假单'
    
    title_run = title_paragraph.add_run(title_text)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 学院称呼
    college_title_paragraph = doc.add_paragraph()
    college_title_run = college_title_paragraph.add_run('各二级学院：')
    college_title_run.font.name = '宋体'
    college_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    college_title_run.font.size = Pt(12)
    college_title_run.font.bold = True

    # 正文内容
    if doc_type == "公假单":
        # 第一段文字
        text_paragraph1 = doc.add_paragraph()
        text_paragraph1.paragraph_format.first_line_indent = Pt(21)
        text_content1 = f'兹定于{activity_info["activity_date"]}举办"{activity_info["activity_name"]}"活动。以下同学因参与活动组织工作，将于{activity_info["work_date"]} {activity_info["work_time"]}协助相关会务工作，无法参加该时间段课程。'
        text_run1 = text_paragraph1.add_run(text_content1)
        text_run1.font.name = '宋体'
        text_run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run1.font.size = Pt(10.5)

        # 第二段文字
        text_paragraph2 = doc.add_paragraph()
        text_paragraph2.paragraph_format.first_line_indent = Pt(21)
        text_content2 = f'特此申请为以下同学办理 {activity_info["work_date"]} {activity_info["work_time"]} 的公假手续，恳请贵学院予以批准，谢谢！'
        text_run2 = text_paragraph2.add_run(text_content2)
        text_run2.font.name = '宋体'
        text_run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run2.font.size = Pt(10.5)
    
    elif doc_type == "抵晚单":
        # 第一段文字
        text_paragraph1 = doc.add_paragraph()
        text_paragraph1.paragraph_format.first_line_indent = Pt(21)
        text_content1 =f'兹定于{activity_info["activity_date"]}举办"{activity_info["activity_name"]}"活动。以下同学因参与活动组织工作，将于{activity_info["work_date"]} {activity_info["work_time"]}协助相关会务工作，无法参加当晚晚自习。'
        text_run1 = text_paragraph1.add_run(text_content1)
        text_run1.font.name = '宋体'
        text_run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run1.font.size = Pt(10.5)

        # 第二段文字
        text_paragraph2 = doc.add_paragraph()
        text_paragraph2.paragraph_format.first_line_indent = Pt(21)
        text_content2 = '特此申请为以下同学办理晚自习请假手续，恳请贵学院予以批准，谢谢！'
        text_run2 = text_paragraph2.add_run(text_content2)
        text_run2.font.name = '宋体'
        text_run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run2.font.size = Pt(10.5)
    
    else:  # 早自习请假单
        # 第一段文字
        text_paragraph1 = doc.add_paragraph()
        text_paragraph1.paragraph_format.first_line_indent = Pt(21)
        text_content1 = f'兹定于{activity_info["activity_date"]}举办"{activity_info["activity_name"]}"活动。以下同学因参与活动组织工作，将于{activity_info["work_date"]} {activity_info["work_time"]}协助相关会务工作，无法参加上午的早自习。'
        text_run1 = text_paragraph1.add_run(text_content1)
        text_run1.font.name = '宋体'
        text_run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run1.font.size = Pt(10.5)

        # 第二段文字
        text_paragraph2 = doc.add_paragraph()
        text_paragraph2.paragraph_format.first_line_indent = Pt(21)
        text_content2 = '特此申请为以下同学办理早自习请假手续，恳请贵学院予以批准，谢谢！'
        text_run2 = text_paragraph2.add_run(text_content2)
        text_run2.font.name = '宋体'
        text_run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text_run2.font.size = Pt(10.5)

    # 创建表格
    table = doc.add_table(rows=1, cols=len(selected_columns))
    table.style = "Table Grid"
    table.autofit=True
    
    # 表头
    header_cells = table.rows[0].cells
    for i, col in enumerate(selected_columns):
        header_cells[i].text = ''
        paragraph = header_cells[i].paragraphs[0]
        run = paragraph.add_run(str(col))
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
        run.font.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(selected_columns):
            value = row[col]
            row_cells[i].text = ''
            paragraph = row_cells[i].paragraphs[0]
            text_content = str(value) if pd.notna(value) else ""
            run = paragraph.add_run(text_content)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 落款
    doc.add_paragraph()
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run1 = signature_paragraph.add_run('共青团温州理工学院委员会')
    run1.font.name = '宋体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run1.font.size = Pt(10.5)
    run1.font.bold = True
    signature_paragraph.add_run('\n')
    
    run2 = signature_paragraph.add_run(activity_info['signature_date'])
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(10.5)
    
    return doc

# 主程序开始
if excel_file is not None:
    try:
        # 读取文件扩展名
        file_extension = excel_file.name.split('.')[-1].lower()
        
        # 智能检测表头位置
        st.info("正在分析Excel文件结构...")
        
        # 预览文件前几行
        engine = 'openpyxl' if file_extension == 'xlsx' else 'xlrd'
        preview_df = pd.read_excel(excel_file, nrows=5, engine=engine)
        excel_file.seek(0)
        
        st.write("**文件前几行预览：**")
        st.dataframe(preview_df)
        
        # 检测表头
        header_row = 0
        for i in range(3):
            row_df = pd.read_excel(excel_file, header=i, nrows=0, engine=engine)
            excel_file.seek(0)
            
            column_names = [str(col).strip().lower() for col in row_df.columns]
            if '学院' in column_names:
                header_row = i
                st.success(f"✅ 检测到表头在第 {header_row + 1} 行")
                break
        
        # 读取完整数据
        df = pd.read_excel(excel_file, header=header_row, engine=engine)
        df.columns = df.columns.str.strip()
        
        # 显示数据预览
        st.write("**数据预览：**")
        st.write(f"总行数：{len(df)}")
        st.dataframe(df.head(10))
        
        # 处理学院列
        college_column = None
        for col in df.columns:
            if '学院' in str(col):
                college_column = col
                break
        
        if college_column is None:
            st.error("未找到包含'学院'的列")
            st.stop()
        
        if college_column != '学院':
            df = df.rename(columns={college_column: '学院'})
        
        # 清理学院列
        df['学院'] = df['学院'].astype(str).str.strip()
        
        # 规范化学院名称
        college_name_mapping = {
            "经管学院": "经济与管理学院",
            "经管": "经济与管理学院",
            "文传学院": "文学与传媒学院",
            "文传": "文学与传媒学院",
            "电电学院": "电子与电气工程学院",
            "电子电气": "电子与电气工程学院",
            "建工学院": "建筑与能源工程学院",
            "建工": "建筑与能源工程学院",
            "外院": "外国语学院",
            "外语": "外国语学院",
            "设艺学院": "设计艺术学院",
            "设计": "设计艺术学院",
            "创业学院": "创新创业学院",
            "数智学院": "数据科学与人工智能学院",
            "数智": "数据科学与人工智能学院",
            "机器人": "机器人工程学院",
            "法学": "法学院"
        }
        
        df["学院"] = df["学院"].apply(lambda x: college_name_mapping.get(str(x).strip(), str(x).strip()))
        
        # 按指定顺序排序
        sorted_dfs = []
        for college in COLLEGE_ORDER:
            college_data = df[df['学院'] == college]
            if not college_data.empty:
                sorted_dfs.append(college_data)
        
        if sorted_dfs:
            df_sorted = pd.concat(sorted_dfs, ignore_index=True)
            other_colleges = set(df['学院'].unique()) - set(COLLEGE_ORDER)
            if other_colleges:
                other_data = df[df['学院'].isin(other_colleges)]
                df_sorted = pd.concat([df_sorted, other_data], ignore_index=True)
            df = df_sorted
        
        # 选择文档类型
        st.write("**第二步：选择请假单类型**")
        doc_type = st.radio(
            "请选择要生成的请假单类型：",
            ["公假单", "抵晚单", "早自习请假单"],
            horizontal=True
        )
        
        # 填写活动信息
        st.write("**第三步：填写活动信息**")
        col1, col2 = st.columns(2)
        
        with col1:
            activity_name = st.text_input("活动名称", "学术讲座")
            work_date = st.text_input("工作日期（如：X月X日）", "X月X日")
            signature_date = st.text_input("落款日期（如：xx年xx月xx日）", "xx年xx月xx日")
        
        with col2:
            if doc_type == "公假单":
                activity_date = st.text_input("活动举办日期（如：X年X月X日）", "X年X月X日")
                work_time = st.selectbox("工作时间段", ["上午", "下午", "全天"])
            else:  # 抵晚单或早自习请假单
                activity_date = ""
                work_time = ""
        
        # 显示填写的活动信息预览
        st.write("**活动信息预览：**")
        st.info(f"请假单类型：{doc_type}")
        st.info(f"活动名称：{activity_name}")
        if doc_type == "公假单":
            st.info(f"活动日期：{activity_date}")
            st.info(f"工作时间：{work_date} {work_time}")
        else:
            st.info(f"请假时间：{work_date}")
        st.info(f"落款日期：{signature_date}")
        
        activity_info = {
            'activity_name': activity_name,
            'activity_date': activity_date,
            'work_date': work_date,
            'work_time': work_time,
            'signature_date': signature_date
        }
        
        # 选择列
        st.write("**第四步：选择要导出的列**")
        all_columns = df.columns.tolist()
        selected_columns = st.multiselect(
            "选择表格中要显示的列",
            all_columns,
            default=all_columns[:min(4, len(all_columns))]
        )
        
        if selected_columns:
            st.write("**表格预览：**")
            st.dataframe(df[selected_columns].head(10))
        
        # 生成文档
        st.write("**第五步：生成文档**")
        if st.button("生成Word文档") and selected_columns:
            with st.spinner("正在生成文档..."):
                doc = create_word_document(df, selected_columns, doc_type, activity_info)
                
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                st.success("✅ 文档生成成功！")
                st.download_button(
                    label="📥 下载Word文档",
                    data=file_stream,
                    file_name=f"{doc_type}_{activity_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    except Exception as e:
        st.error(f"处理文件失败: {str(e)}")

else:
    st.info("请先上传Excel文件（支持.xlsx和.xls格式）")

